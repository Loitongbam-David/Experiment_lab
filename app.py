import streamlit as st
from openai import OpenAI
import base64
import time
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
import io
import re
import PyPDF2

# --- CONFIGURATION ---
st.set_page_config(page_title="Universal Lab Agent (Smart Analyzer)", layout="wide")

# --- HARDCODED SETTINGS ---
TEMPLATE_PATH = "sample.docx"

# --- TRANSLATIONS DICTIONARY ---
TRANS = {
    "English": {
        "title": "Universal Lab Agent & Tutor ",
        "tab_gen": " Experiment Generator",
        "tab_analyze": "📂 Document Analyzer",
        "sidebar_config": "⚙️ Configuration",
        "lbl_language": "Programming Language",
        "lbl_ui_lang": "Interface Language / भाषा",
        "lbl_upload": "1. Upload Problem Statement (Image)",
        "lbl_drag_drop": " Drag & Drop Image (Diagram/Screenshot)",
        "lbl_details": "2. Enter Experiment Details",
        "lbl_exp_name": "Experiment Name / Topic",
        "lbl_exp_ph": "e.g. Implement BFS...",
        "btn_generate": " Generate Record",
        "warn_input": "Please enter a topic OR upload an image.",
        "spinner_gen": "Generating Experiment Record...",
        "spinner_analyze": "Isolating experiment logic & Analyzing...",
        "success": "✅ Success!",
        "download": "📥 Download Word File",
        "header_explain": " Detailed Breakdown",
        "header_viva": " Viva Preparation",
        "expander_viva": "Show Viva Questions & Answers",
        "header_chat": "💬 Chat with AI Tutor",
        "chat_placeholder": "Ask a doubt about this experiment...",
        "lbl_doc_upload": "Upload Experiment Doc (PDF/DOCX)",
        "btn_analyze": "🔍 Analyze Document",
        "warn_doc": "Please upload a document first.",
        "err_empty": "AI returned empty data. Try again."
    },
    "Hindi": {
        "title": "यूनिवर्सल लैब एजेंट और ट्यूटर ",
        "tab_gen": " प्रयोग (Experiment) जनरेटर",
        "tab_analyze": "📂 दस्तावेज़ (Document) एनालाइज़र",
        "sidebar_config": "⚙️ सेटिंग्स",
        "lbl_language": "प्रोग्रामिंग भाषा",
        "lbl_ui_lang": "Interface Language / भाषा",
        "lbl_upload": "1. समस्या की छवि अपलोड करें",
        "lbl_drag_drop": " चित्र यहाँ ड्रॉप करें",
        "lbl_details": "2. प्रयोग विवरण",
        "lbl_exp_name": "प्रयोग का नाम / विषय",
        "lbl_exp_ph": "उदा. BFS Implement करें...",
        "btn_generate": " रिकॉर्ड जनरेट करें",
        "warn_input": "कृपया विषय लिखें या इमेज दें।",
        "spinner_gen": "रिकॉर्ड तैयार किया जा रहा है...",
        "spinner_analyze": "प्रयोग को समझा जा रहा है...",
        "success": "✅ सफल!",
        "download": "📥 वर्ड फाइल डाउनलोड करें",
        "header_explain": " विस्तृत स्पष्टीकरण (Explanation)",
        "header_viva": " वाइवा (Viva) तैयारी",
        "expander_viva": "वाइवा प्रश्न और उत्तर देखें",
        "header_chat": "💬 AI ट्यूटर से चैट करें",
        "chat_placeholder": "इस प्रयोग के बारे में पूछें...",
        "lbl_doc_upload": "प्रयोग दस्तावेज़ अपलोड करें (PDF/DOCX)",
        "btn_analyze": "🔍 विश्लेषण (Analyze) करें",
        "warn_doc": "कृपया पहले दस्तावेज़ अपलोड करें।",
        "err_empty": "डेटा नहीं मिला। पुनः प्रयास करें।"
    }
}

# --- SESSION STATE ---
if "doc_buffer" not in st.session_state: st.session_state.doc_buffer = None
if "experiment_name" not in st.session_state: st.session_state.experiment_name = ""
if "explanation_text" not in st.session_state: st.session_state.explanation_text = ""
if "code_text" not in st.session_state: st.session_state.code_text = ""
if "viva_questions" not in st.session_state: st.session_state.viva_questions = ""
if "chat_history" not in st.session_state: st.session_state.chat_history = []
if "uploaded_image_bytes" not in st.session_state: st.session_state.uploaded_image_bytes = None
if "ui_lang" not in st.session_state: st.session_state.ui_lang = "English"

# --- AI SETUP ---
try:
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except FileNotFoundError:
    GROQ_API_KEY = "" 

client = OpenAI(base_url="https://api.groq.com/openai/v1", api_key=GROQ_API_KEY)

# --- HELPER FUNCTIONS ---

def get_txt(key):
    return TRANS[st.session_state.ui_lang][key]

def get_image_base64_and_mime(image_file):
    if image_file:
        return base64.b64encode(image_file.getvalue()).decode('utf-8'), image_file.type
    return None, None

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def filter_metadata(text):
    """
    Scans text for the start of the actual experiment (Aim, Problem, Code)
    and removes the student metadata (Name, UID, Branch) found before it.
    """
    lower_text = text.lower()
    # Keywords that usually signify the start of the experiment content
    keywords = ["aim", "aim:", "objective", "problem statement", "experiment no", "concept:", "code:", "program:"]
    
    start_index = 0
    found_keyword = False
    
    for kw in keywords:
        idx = lower_text.find(kw)
        if idx != -1:
            # We found a keyword. To be safe, we back up a few chars (e.g., "1. Aim")
            # But generally, cutting from here is safer than keeping metadata.
            start_index = idx
            found_keyword = True
            break
            
    if found_keyword:
        return text[start_index:] # Return only text AFTER metadata
    return text # Return full text if no keywords found

def format_viva_output(text):
    formatted = re.sub(r'(\d+\.)', r'\n\n\1', text)
    formatted = re.sub(r'\n{3,}', r'\n\n', formatted)
    return formatted.strip()

def format_explanation_output(text):
    formatted = re.sub(r'(Lines? \d+)', r'\n- **\1', text)
    formatted = formatted.replace("- **- **", "- **") 
    formatted = re.sub(r'\n{3,}', r'\n\n', formatted)
    return formatted.strip()

def get_ai_response(prompt, image_data=None):
    try:
        if image_data and image_data[0]:
            b64_str, mime_type = image_data
            response = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct", 
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{b64_str}"}},
                        ],
                    }
                ],
                temperature=0.5,
                max_completion_tokens=2048 
            )
        else:
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.6
            )
        return response.choices[0].message.content
    except Exception as e:
        return f"ERROR_DETAILS: {str(e)}"

def insert_text_after_paragraph(paragraph, text, bold=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    if bold:
        run = new_para.add_run(text)
        run.bold = True
    else:
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if line.strip():
                if i > 0: 
                    new_para = paragraph._parent.add_paragraph(line)
                    paragraph._p.addnext(new_para._p)
                    paragraph = new_para
                else:
                    new_para.add_run(line)
    paragraph._p.addnext(new_para._p)
    return new_para

def fill_template(template_path, content_dict, image_bytes=None):
    doc = Document(template_path)
    mapping = {
        "Aim/Overview": "Aim", "Concept/Algorithm": "Concept", 
        "Coding": "Code", "Screenshot and Output": "Output", "Learning outcomes": "Conclusion"
    }
    for para in doc.paragraphs:
        text = para.text.lower().strip()
        for doc_key, ai_key in mapping.items():
            if doc_key.lower() in text and len(text) < 100:
                if ai_key == "Output":
                    if "Execution Steps" in content_dict and content_dict.get("Execution Steps"):
                        insert_text_after_paragraph(para, content_dict["Execution Steps"])
                    if content_dict.get("Output"):
                        insert_text_after_paragraph(para, "\nSample Output:", bold=True)
                        insert_text_after_paragraph(para, content_dict["Output"])
                    if image_bytes:
                        p = doc.add_paragraph()
                        para._p.addnext(p._p)
                        p.add_run("\nScreenshot / Diagram:\n").add_picture(io.BytesIO(image_bytes), width=Inches(4.0))
                    content_dict["Output"] = None
                elif ai_key in content_dict and content_dict[ai_key]:
                    val = content_dict[ai_key]
                    if ai_key == "Code":
                        for line in reversed(val.split('\n')):
                            p = insert_text_after_paragraph(para, line)
                            p.paragraph_format.space_after = Pt(0)
                            if p.runs: 
                                p.runs[0].font.name = 'Courier New'
                                p.runs[0].font.size = Pt(10)
                    else:
                        insert_text_after_paragraph(para, val)
                    content_dict[ai_key] = None
                break
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def clean_and_parse(raw_text):
    text = re.sub(r'<think>.*?</think>', '', raw_text, flags=re.DOTALL)
    content = {}
    current_section = None
    valid_keys = ["Aim", "Concept", "Code", "Output", "Conclusion", "Explanation", "Viva Questions", "Execution Steps"]
    for line in text.split('\n'):
        clean_line = line.strip().replace('#', '').replace(':', '').strip()
        is_header = False
        for key in valid_keys:
            if key.lower() == clean_line.lower():
                current_section = key
                content[key] = ""
                is_header = True
                break
        if not is_header and current_section:
            if "```" not in line: content[current_section] += line + "\n"
    
    if "Viva Questions" in content: content["Viva Questions"] = format_viva_output(content["Viva Questions"])
    if "Explanation" in content: content["Explanation"] = format_explanation_output(content["Explanation"])
    return content

# --- UI START ---

with st.sidebar:
    st.header(TRANS["English"]["sidebar_config"])
    st.session_state.ui_lang = st.radio(get_txt("lbl_ui_lang"), ["English", "Hindi"], horizontal=True)
    st.divider()
    programming_language = st.selectbox(get_txt("lbl_language"), ["Python", "C++","C", "Java", "SQL"])

st.title(get_txt("title"))

tab_gen, tab_analyze = st.tabs([get_txt("tab_gen"), get_txt("tab_analyze")])

# ==========================================
# TAB 1: EXPERIMENT GENERATOR
# ==========================================
with tab_gen:
    st.subheader(get_txt("lbl_upload"))
    uploaded_file = st.file_uploader(get_txt("lbl_drag_drop"), type=['png', 'jpg', 'jpeg'], key="gen_img")
    
    if uploaded_file:
        with st.expander("Preview", expanded=False): st.image(uploaded_file, width=300)
        st.session_state.uploaded_image_bytes = uploaded_file.getvalue()
    
    st.divider()
    st.subheader(get_txt("lbl_details"))
    
    input_height = 68 
    if st.session_state.experiment_name:
        num_lines = st.session_state.experiment_name.count('\n') + (len(st.session_state.experiment_name) // 80)
        input_height = min(300, max(68, num_lines * 30))

    experiment_topic = st.text_area(get_txt("lbl_exp_name"), value=st.session_state.experiment_name, height=input_height, key="gen_input")

    if st.button(get_txt("btn_generate"), type="primary", key="btn_gen"):
        st.session_state.experiment_name = experiment_topic
        if not experiment_topic and not uploaded_file:
            st.warning(get_txt("warn_input"))
        else:
            with st.spinner(get_txt("spinner_gen")):
                image_data = get_image_base64_and_mime(uploaded_file)
                topic_str = f"Topic: {experiment_topic}" if experiment_topic else "Topic: Analyze image & solve."
                lang_rule = "Use Hinglish for Explanation/Viva." if st.session_state.ui_lang == "Hindi" else ""
                
                prompt = f"""
                Act as Lab Instructor. {topic_str}. Lang: {programming_language}. {lang_rule}
                Provide 8 sections: ### Aim, ### Concept, ### Code, ### Execution Steps, ### Output, ### Conclusion, ### Explanation, ### Viva Questions.
                Rules: 
                - Code raw only. 
                - Explanation MUST be a Bulleted list.
                - Viva 25 Qs (Basic,Logic,Advanced) with answers.
                """
                
                raw = get_ai_response(prompt, image_data)
                parsed = clean_and_parse(raw)
                
                if parsed:
                    if "Code" in parsed: st.session_state.code_text = parsed["Code"]
                    if "Explanation" in parsed: st.session_state.explanation_text = parsed["Explanation"]
                    if "Viva Questions" in parsed: st.session_state.viva_questions = parsed["Viva Questions"]
                    if "Aim" in parsed: st.session_state.experiment_name = parsed["Aim"][:200]
                    
                    st.session_state.chat_history = [] 
                    st.session_state.doc_buffer = fill_template(TEMPLATE_PATH, parsed, st.session_state.uploaded_image_bytes)
                    st.success(get_txt("success"))
                    st.rerun()

    if st.session_state.doc_buffer:
        st.download_button(get_txt("download"), st.session_state.doc_buffer, "Experiment.docx")

# ==========================================
# TAB 2: DOCUMENT ANALYZER (Updated)
# ==========================================
with tab_analyze:
    st.info("Upload your existing experiment file and get your experiment explained in depth")
    
    doc_file = st.file_uploader(get_txt("lbl_doc_upload"), type=['pdf', 'docx'], key="doc_up")

    st.markdown("---- or -----")

    manual_input=st.text_area("Enter code or experiment to be explained")
    if st.button(get_txt("btn_analyze"), type="primary", key="btn_an"):
        if not doc_file and not manual_input.strip():
            st.warning(get_txt("warn_doc"))
        else:
            with st.spinner(get_txt("spinner_analyze")):
              if doc_file:
                if doc_file.type == "application/pdf":
                    full_text = extract_text_from_pdf(doc_file)
                else:
                    full_text = extract_text_from_docx(doc_file)               
                # --- APPLY SMART FILTER ---
                clean_text = filter_metadata(full_text)
              else:
                clean_text=manual_input
              
              st.session_state.code_text = clean_text[:3000] # Store for chat context
              
              lang_rule = "Write Explanation and Viva in HINDI (Hinglish)." if st.session_state.ui_lang == "Hindi" else "Write in English."
              
              prompt = f"""
              Analyze the following technical content (Code/Concept):
              {clean_text[:6000]} 
              
              {lang_rule}
              
              Provide 2 sections EXACTLY:
              ### Explanation
              ### Viva Questions
              
              Rules:
              1. **Explanation:** Ignore the 'Aim' text itself. Focus strictly on explaining the **Concept Steps** and **Line-by-Line Code Logic**. 
                 - Use bullet points (e.g., "- **Step 1:**...", "- **Lines 10-12:**...").
              2. **Viva Questions:** Generate exactly 25 Questions (10 Basic, 10 Logic, 5 Advanced) WITH ANSWERS based on this code.
              """
              
              raw = get_ai_response(prompt)
              parsed = clean_and_parse(raw)
              
              if parsed:
                  st.session_state.explanation_text = parsed.get("Explanation", "No explanation found.")
                  st.session_state.viva_questions = parsed.get("Viva Questions", "No questions found.")
                  st.session_state.chat_history = [] 
                  st.success(get_txt("success"))
              else:
                  st.error(get_txt("err_empty"))

# ==========================================
# SHARED OUTPUT SECTION
# ==========================================

if st.session_state.explanation_text:
    st.divider()
    st.subheader(get_txt("header_explain"))
    with st.container(border=True):
        st.markdown(st.session_state.explanation_text)

if st.session_state.viva_questions:
    st.divider()
    st.subheader(get_txt("header_viva"))
    with st.expander(get_txt("expander_viva"), expanded=True):
        st.markdown(st.session_state.viva_questions)

st.divider()
st.subheader(get_txt("header_chat"))

for msg in st.session_state.chat_history:
    with st.chat_message(msg["role"]): st.markdown(msg["content"])

if user_input := st.chat_input(get_txt("chat_placeholder")):
    st.session_state.chat_history.append({"role": "user", "content": user_input})
    with st.chat_message("user"): st.markdown(user_input)
    
    lang_pref = "Reply in Hindi/Hinglish." if st.session_state.ui_lang == "Hindi" else "Reply in English."
    
    ctx_prompt = f"""
    Act as Tutor. Lang: {lang_pref}.
    CONTEXT: {st.session_state.code_text[:3000]}...
    EXPLANATION: {st.session_state.explanation_text[:2000]}...
    QUESTION: {user_input}
    Keep answer short & clear.
    """
    
    with st.chat_message("assistant"):
        with st.spinner("..."):
            reply = get_ai_response(ctx_prompt)
            st.markdown(reply)
            st.session_state.chat_history.append({"role": "assistant", "content": reply})